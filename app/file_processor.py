# app/file_processor.py
import os
import uuid
import aiofiles
from typing import List, Dict, Any
from PyPDF2 import PdfReader
import pdfplumber
from docx import Document
import chardet

class FileProcessor:
    """Process uploaded files for PDF, DOCX, and TXT formats."""
    
    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
    
    def __init__(self, upload_dir: str = "./uploads"):
        self.upload_dir = upload_dir
        os.makedirs(upload_dir, exist_ok=True)
    
    async def save_file(self, file) -> str:
        """Save the uploaded file and return its path."""
        file_id = str(uuid.uuid4())
        original_name = file.filename
        extension = os.path.splitext(original_name)[1].lower()
        
        if extension not in self.ALLOWED_EXTENSIONS:
            raise ValueError(f"File type not supported. Supported types: {self.ALLOWED_EXTENSIONS}")
        
        new_filename = f"{file_id}{extension}"
        file_path = os.path.join(self.upload_dir, new_filename)
        
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        
        return file_path, file_id, original_name
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from file based on its type."""
        extension = os.path.splitext(file_path)[1].lower()
        
        if extension == '.pdf':
            return self._extract_from_pdf(file_path)
        elif extension == '.docx':
            return self._extract_from_docx(file_path)
        elif extension == '.txt':
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"File type not supported: {extension}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF - dual method for accuracy."""
        text = ""
        
        # First attempt using pdfplumber (better for tables)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except:
            pass
        
        # If failed, use PyPDF2
        if not text.strip():
            try:
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except:
                pass
        
        return text.strip() if text.strip() else "Failed to extract text from PDF file"
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # Extract text from tables as well
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += "\n" + cell.text.strip()
            
            return text.strip() if text.strip() else "No text found in the file"
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file with encoding detection."""
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                detected = chardet.detect(raw_data)
                encoding = detected.get('encoding', 'utf-8')
            
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except Exception as e:
            return f"Error reading TXT: {str(e)}"
    
    def delete_file(self, file_path: str):
        """Delete the file after processing."""
        if os.path.exists(file_path):
            os.remove(file_path)